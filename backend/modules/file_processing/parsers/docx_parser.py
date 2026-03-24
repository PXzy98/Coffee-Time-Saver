import io
from docx import Document
from modules.file_processing.parsers.base import FileParser


class DocxParser(FileParser):
    async def parse(self, file_bytes: bytes, filename: str) -> str:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append("\t".join(cells))
        return "\n\n".join(paragraphs)
