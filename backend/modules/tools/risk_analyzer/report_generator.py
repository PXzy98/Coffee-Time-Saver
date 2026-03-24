"""Generate downloadable PDF or DOCX report from a RiskReport."""
import io
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

from modules.tools.risk_analyzer.schemas import RiskReport


def generate_pdf(report: RiskReport) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("Project Risk Analysis Report", styles["Title"]))
    story.append(Paragraph(f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}", styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    # Executive Summary
    story.append(Paragraph("Executive Summary", styles["Heading1"]))
    story.append(Paragraph(f"Overall Risk Level: <b>{report.overall_risk_level.upper()}</b>", styles["Normal"]))
    story.append(Paragraph(f"Overall Confidence: {report.overall_confidence:.0%}", styles["Normal"]))
    story.append(Paragraph(report.executive_summary, styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    # Risk Register
    if report.risks:
        story.append(Paragraph("Risk Register", styles["Heading1"]))
        data = [["#", "Description", "Category", "Likelihood", "Impact", "Score", "Conf."]]
        for i, risk in enumerate(report.risks, 1):
            data.append([
                str(i),
                Paragraph(risk.description[:120], styles["Normal"]),
                risk.category,
                str(risk.likelihood),
                str(risk.impact),
                f"{risk.risk_score:.2f}",
                f"{risk.confidence:.0%}",
            ])
        table = Table(data, colWidths=[0.8*cm, 5.5*cm, 2*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.5*cm))

        for risk in report.risks:
            story.append(Paragraph(f"<b>Mitigation ({risk.id}):</b> {risk.mitigation}", styles["Normal"]))
        story.append(Spacer(1, 0.5*cm))

    # Inconsistencies
    if report.inconsistencies:
        story.append(Paragraph("Inconsistency Report", styles["Heading1"]))
        for item in report.inconsistencies:
            story.append(Paragraph(f"<b>[{item.type.upper()}]</b> {item.explanation}", styles["Normal"]))
            story.append(Paragraph(f"<i>{item.document_a}:</i> {item.passage_a[:200]}", styles["Normal"]))
            story.append(Paragraph(f"<i>{item.document_b}:</i> {item.passage_b[:200]}", styles["Normal"]))
            story.append(Paragraph(f"Recommendation: {item.recommendation}", styles["Normal"]))
            story.append(Spacer(1, 0.3*cm))

    # Appendix
    story.append(Paragraph("Appendix", styles["Heading1"]))
    story.append(Paragraph("Documents Analyzed:", styles["Heading2"]))
    for fname in report.documents_analyzed:
        story.append(Paragraph(f"• {fname}", styles["Normal"]))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("Methodology:", styles["Heading2"]))
    story.append(Paragraph(report.methodology_notes, styles["Normal"]))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(report: RiskReport) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    import io

    doc = Document()
    doc.add_heading("Project Risk Analysis Report", 0)
    doc.add_paragraph(f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Overall Risk Level: {report.overall_risk_level.upper()}")
    doc.add_paragraph(f"Overall Confidence: {report.overall_confidence:.0%}")
    doc.add_paragraph(report.executive_summary)

    if report.risks:
        doc.add_heading("Risk Register", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Description", "Category", "Likelihood", "Impact", "Score", "Confidence"]):
            hdr[i].text = h
        for risk in report.risks:
            row = table.add_row().cells
            row[0].text = risk.description[:150]
            row[1].text = risk.category
            row[2].text = str(risk.likelihood)
            row[3].text = str(risk.impact)
            row[4].text = f"{risk.risk_score:.2f}"
            row[5].text = f"{risk.confidence:.0%}"
        doc.add_paragraph("")
        for risk in report.risks:
            doc.add_paragraph(f"Mitigation ({risk.id}): {risk.mitigation}")

    if report.inconsistencies:
        doc.add_heading("Inconsistency Report", level=1)
        for item in report.inconsistencies:
            doc.add_paragraph(f"[{item.type.upper()}] {item.explanation}")
            doc.add_paragraph(f"{item.document_a}: {item.passage_a[:200]}")
            doc.add_paragraph(f"{item.document_b}: {item.passage_b[:200]}")
            doc.add_paragraph(f"Recommendation: {item.recommendation}")

    doc.add_heading("Appendix", level=1)
    doc.add_heading("Documents Analyzed", level=2)
    for fname in report.documents_analyzed:
        doc.add_paragraph(f"• {fname}")
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(report.methodology_notes)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
